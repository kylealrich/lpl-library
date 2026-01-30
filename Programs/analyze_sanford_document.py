import docx
import json
from pathlib import Path

def analyze_document():
    doc_path = Path(r"c:\lpl-library\Inputs\Sanford_ANA-050+DES-020_INT_003_Purchase Order Receipt Inbound Interface.docx")
    
    doc = docx.Document(doc_path)
    
    analysis = {
        "document_title": "Sanford Purchase Order Receipt Inbound Interface",
        "total_paragraphs": len(doc.paragraphs),
        "total_tables": len(doc.tables),
        "content_sections": [],
        "tables_data": [],
        "key_information": {}
    }
    
    # Extract paragraph content
    for i, para in enumerate(doc.paragraphs):
        if para.text.strip():
            analysis["content_sections"].append({
                "paragraph": i + 1,
                "text": para.text.strip()
            })
    
    # Extract table data
    for i, table in enumerate(doc.tables):
        table_data = []
        for row in table.rows:
            row_data = [cell.text.strip() for cell in row.cells]
            table_data.append(row_data)
        analysis["tables_data"].append({
            "table": i + 1,
            "rows": len(table.rows),
            "columns": len(table.columns) if table.rows else 0,
            "data": table_data
        })
    
    # Save analysis to output
    output_path = Path(r"c:\lpl-library\Outputs\sanford_document_analysis.json")
    with open(output_path, 'w', encoding='utf-8') as f:
        json.dump(analysis, f, indent=2, ensure_ascii=False)
    
    print(f"Document Analysis Complete:")
    print(f"- Paragraphs: {analysis['total_paragraphs']}")
    print(f"- Tables: {analysis['total_tables']}")
    print(f"- Analysis saved to: {output_path}")
    
    return analysis

if __name__ == "__main__":
    analyze_document()